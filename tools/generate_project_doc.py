from docx import Document
from docx.shared import Pt
from datetime import datetime

# Comprehensive project documentation content
TITLE = "Pet Adoption & Rescue Management Portal - Project Documentation"
FILENAME = "PROJECT_DOCUMENTATION.docx"

content = {
    "Overview": (
        "This document describes the Pet Adoption & Rescue Management Portal project. "
        "It summarizes architecture, main features, data model, multi-database chat design, developer "
        "setup, testing and deployment notes, and next steps."
    ),

    "Repository & Branch": (
        "Repository: Pet-Adoption-and-Rescue-Management-Portal_September_Batch-3_2025\n"
        "Current branch (local): Abhijit_Singh_local_push\n"
        "Note: local feature branch was pushed as Abhijit_Singh_local_push to preserve work when the original branch "
        "Abhijit_Singh on remote diverged. A PR URL was created from that branch."
    ),

    "Tech Stack": (
        "- Backend: Django 5.2.x (Python 3.13.x)\n"
        "- Database: Primary DB (MySQL in production / sqlite for local dev), plus a separate chat DB (sqlite) used for chat models\n"
        "- Templates: Django templates (project ships both legacy and 'modern' templates)\n"
        "- Frontend: Minimal JS for chat polling / send via fetch; CSS under webapp/static/webapp\n"
    ),

    "High-level Architecture": (
        "- Single Django project using two DB connections: the default DB for main models (users, pets, admin data, notifications), and a secondary sqlite DB alias 'chat_db' for chat-related tables.\n"
        "- Chat models are deliberately decoupled from auth.User to avoid cross-database foreign key constraints: messages and chat membership store integer user ids (sender_id, user_id).\n"
        "- Views that operate on chat data use .using('chat_db') for writes and carefully avoid ORM operations that would join across databases.\n"
    ),

    "Key Features": (
        "- Pet listings (add/edit/list, statuses for adoption/lost/found)\n"
        "- Adoption request flow and notifications to pet owners\n"
        "- In-app Chat saved in separate database (chat_db) with Conversation, Message, ChatMember models\n"
        "- Notifications model and context processor for navbar badge counts\n"
        "- Admin features: start chat with any user, admin dashboard, admin-only message deletion, conversation deletion\n"
        "- Legal pages: About, Contact, Privacy, Terms and small UI improvements (responsive navbar, home punchline).\n"
    ),

    "Chat Design Details": (
        "- chat/models.py contains Conversation, Message, ChatMember (managed on chat_db), and an unmanaged ChatParticipant mapping for legacy through tables.\n"
        "- Messages store sender_id (integer) and conversation_id (FK to Conversation defined in chat DB).\n"
        "- Chat writes and deletes use the 'chat_db' alias; some delete operations use raw SQL to avoid cascading into unmanaged tables that may not exist.\n"
        "- When creating conversations about pets, conversation.subject is tagged with (pet:<id>) so that starting a chat about a different pet creates a new conversation (avoids reusing the same convo for multiple pets).\n"
    ),

    "Important Files and Locations": (
        "- manage.py (project runner)\n"
        "- home/ (project settings and WSGI/ASGI)\n"
        "- webapp/ (main app): models.py, views.py, templates/webapp/, static/webapp/\n"
        "- chat/ (chat app): models.py, views.py, templates/chat/, migrations/, scripts/\n"
        "- db.sqlite3 (local default DB), chat_db.sqlite3 (chat DB at project root)\n"
    ),

    "Developer Setup (quick)": (
        "1. Create and activate a Python 3.13 venv.\n"
        "2. Install requirements (project does not ship a pinned requirements.txt). Minimum: Django 5.2.x.\n"
        "3. Configure your local settings (home/settings.py) if you need to point to MySQL or use local sqlite.\n"
        "4. Run migrations for default DB: python manage.py migrate\n"
        "5. Run chat DB migrations specifically: python manage.py migrate --database=chat_db\n"
        "6. Create a superuser: python manage.py createsuperuser\n"
        "7. Start dev server: python manage.py runserver\n"
    ),

    "Tests & QA": (
        "- Tests for chat are located under chat/tests*.py.\n"
        "- Example test run used in development: python manage.py test chat -v2 --keepdb\n"
        "- The project includes a chat flow script at chat/scripts/flow_run.py used to smoke-test chat flows against a running devserver.\n"
    ),

    "Git / Branching Notes": (
        "- Original feature branch 'Abhijit_Singh' diverged on remote; to avoid force-pushing, local changes were pushed to a new branch 'Abhijit_Singh_local_push'. A PR URL is available for that branch.\n"
        "- Repository history included accidental commits of compiled files in __pycache__; these were removed from the index and should be added to .gitignore to prevent future conflicts.\n"
    ),

    "Known Issues & Recommendations": (
        "- Cross-DB operations are fragile: avoid ORM joins across databases (e.g., do not use select_related on Message -> sender User). The current approach uses integer ids and separate lookups.\n"
        "- Legacy unmanaged join/through tables and cascade deletes required raw SQL in some cases. Keep that in mind when refactoring models or migrations.\n"
        "- Consider adding a small README and a pinned requirements.txt (pip freeze > requirements.txt) for reproducible setups.\n"
    ),

    "API & Next Steps": (
        "- A DRF-based API was proposed: endpoints for Pets, Users, Conversations, Messages.\n"
        "- Note: Chat API must carefully use .using('chat_db') for writes and avoid cross-db FKs; serializers should map sender_id to read-only fields representing usernames on the main DB.\n"
    ),

    "Contact & Maintainers": (
        "- Primary contributor branch: Abhijit_Singh_local_push\n"
        "- For feature questions, open an issue or PR on the repository.\n"
    ),
}


def add_heading_paragraph(doc, heading, text):
    doc.add_heading(heading, level=2)
    for line in text.split('\n'):
        if not line.strip():
            continue
        p = doc.add_paragraph()
        p.style.font.size = Pt(11)
        p.add_run(line)


def main():
    doc = Document()
    doc.core_properties.title = TITLE
    doc.add_heading(TITLE, level=1)
    doc.add_paragraph(f'Generated: {datetime.utcnow().isoformat()} UTC')
    doc.add_paragraph('')

    for section, text in content.items():
        add_heading_paragraph(doc, section, text)
        doc.add_paragraph('')

    # Save
    out_path = FILENAME
    doc.save(out_path)
    print(f'Wrote {out_path}')

if __name__ == '__main__':
    main()
